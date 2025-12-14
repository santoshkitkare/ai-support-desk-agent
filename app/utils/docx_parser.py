##############################################################################
# File: 🚀 docx_parser.py — Your Word Document Extractor
# It plays a big role in making your ingestion service robust and enterprise-ready. 
# DOCX files (especially support SOPs, manuals, HR policies, troubleshooting guides) are 
# extremely common in real companies.
# 
# 🎯 Why This Parser Is More Than Enough for Your Use Case
# Support documents often come in DOCX formats like:
# - warranty guidelines
# - onboarding instructions
# - compliance procedures
# - operational SOPs
# - FAQs exported from internal systems
# They’re usually:
# - paragraph-based
# - not heavily styled
# - not filled with tables
# This parser captures exactly the text you need — nothing more, nothing less.
from docx import Document as DocxDocument


def extract_text_from_docx(path: str) -> str:
    # 1️⃣ Load the DOCX file
    # Uses python-docx to open the document.
    # This library correctly handles:
    # - paragraphs
    # - runs
    # - headings
    # - bold/italic (ignored, but text preserved)
    # - hidden characters
    # - unicode
    doc = DocxDocument(path)
    
    # 2️⃣ Extract text from each paragraph
    # ✔ Why this line is good:
    # - .paragraphs iterates all logical text blocks
    # - .text extracts the plain text
    # - if p.text.strip() filters out empty or whitespace-only lines
    # - Avoids generating tiny, meaningless chunks
    
    # 3️⃣ Combine paragraphs into a single text blob
    # ✔ Why this works well:
    # Easy for chunking
    # Preserves natural paragraph breaks
    # Works across platforms
    # Avoids noise
    paragraphs = [p.text for p in doc.paragraphs if p.text.strip()]
    return "\n".join(paragraphs)


"""
⚠️ Limitations (only if documents are complex)
This is just for awareness — not required for your current setup:

1️⃣ DOCX tables
python-docx doesn’t include table text via .paragraphs.
If your documents have tables like:
Issue	Resolution	Time
You’d need:
    for table in doc.tables:
        for row in table.rows:
            cells = [cell.text for cell in row.cells]
But again — not needed unless your support documents rely heavily on tables.

2️⃣ Headers/Footers are not included
Also fine — usually irrelevant for RAG.

3️⃣ Lists (unordered/bulleted) appear as plain text
That’s acceptable for embeddings.

💡 Optional Enhancement (Not Required)
If you want to get fancier later:
Add structural tags:
    [Heading] Warranty Information  
    Paragraph text...

This can help reranking and improve semantic retrieval.
But again — optional.
"""